# Bài tập tổng hợp phần Basic - Đếm từ
# Hãy vào đường link dưới đây để tải về file .py có viết sẵn khung code và yêu cầu cho bài tập này.
# Có 2 file text là sample input. Bài này sẽ cần dùng 1 chút kiến thức về việc open 1 file. Bạn có thể xem ở notebook buổi sau ở đây hoặc tìm tài liệu trên mạng.
# 2 Câu lệnh ví dụ sau dùng để chạy chương trình khi bạn đã code xong.
# python wordcount.py --count small.txt
# python wordcount.py --topcount small.txt
"""Wordcount exercise

Hàm main() đã được định nghĩa hoàn chỉnh ở dưới. Bạn phải viết hàm get_words()
và get_top_words() mà sẽ được gọi từ main().

1. Với đối số --count, viết hàm get_words(filename) đếm số lần xuất hiện của mỗi từ 
trong file đầu vào và trả list các tuple theo định dạng sau:
[(word1, count1), 
(word2, count2)
...]

Trả ra danh sách trên theo thứ tự từ điển các từ (python sẽ sắp xếp dấu câu đứng trước
các chữ cái nên cũng không thành vấn đề). Lưu tất cả các từ dưới dạng chữ thường,
vì vậy 'The' và 'the' được tính là cùng một từ.

2. Với đối số --topcount, viết hàm get_top_words(filename) tương tự như get_words()
nhưng chỉ trả ra 20 từ thông dụng nhất sắp xếp theo từ thông dụng nhất ở trên cùng.

Tùy chọn: định nghĩa một hàm helper để tránh lặp lại code trong các hàm 
get_words() và get_top_words().

"""
# pip install python-docx
import sys
import docx
import re

# +++your code here+++


def word_count(string, words=dict()):
    my_list = re.split(r'[^a-zA-Z0-9_]', string)
    my_list = [x.lower() for x in my_list]
    unique_word = set(my_list)
    unique_word.discard('')
    for k in unique_word:
        number_count = my_list.count(k)
        if k in words.keys():
            number_count += words.get(k)
        words[k] = number_count

def get_words(filename):
    document = docx.Document(filename)
    words = dict()
    for i in document.paragraphs:
        word_count(i.text, words)
    return sorted(words.items())

def get_top_words(filename):
    document = docx.Document(filename)
    words = dict()
    for i in document.paragraphs:
        word_count(i.text, words)
    words = sorted(words.items(), key=lambda x: x[1], reverse=True)
    return words[0:20]

###
# This basic command line argument parsing code is provided and
# calls the get_words() and get_top_words() functions which you must define.

def main():
    if len(sys.argv) != 3:
        print('usage: ./wordcount.py {--count | --topcount} file')
        sys.exit(1)
    
    option = sys.argv[1]
    filename = sys.argv[2]
    ans = []
    print(filename)
    if option == '--count':
        ans = get_words(filename)
    elif option == '--topcount':
        ans = get_top_words(filename)
    else:
        print('unknown option: ' + option)
        sys.exit(1)
    # print out the answer
    for item in ans:
        print(item[0], item[1])


if __name__ == '__main__':
    main()
